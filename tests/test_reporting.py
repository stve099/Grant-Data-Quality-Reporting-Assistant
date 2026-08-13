"""Report generation and export tests (HTML, Word, Excel)."""

from __future__ import annotations

import openpyxl
import pytest
from docx import Document

from grant_assistant.reporting import (
    build_report_data,
    render_html_report,
    write_analytics_workbook,
    write_audit_workbook,
    write_docx_report,
    write_html_report,
)


@pytest.fixture(scope="module")
def report_data(analytics_flawed, audit_flawed, profile):
    return build_report_data(analytics_flawed, audit_flawed, profile)


def test_report_data_assembled(report_data, analytics_flawed):
    assert report_data.executive_summary
    assert str(analytics_flawed.total_enrollments) in report_data.executive_summary
    assert report_data.insights.key_findings
    assert report_data.ai_generated_narrative is False
    assert report_data.measure_definitions()
    assert report_data.data_limitations()


def test_html_report_contains_all_sections(report_data):
    html = render_html_report(report_data)
    for heading in (
        "Program Overview",
        "Executive Summary",
        "Data Quality Statement",
        "Population Served",
        "Demographic Summary",
        "Enrollment &amp; Exit Metrics",
        "Housing Outcomes",
        "Income Outcomes",
        "Follow-Up Outcomes",
        "Performance Measures",
        "Program Comparison",
        "Key Findings",
        "Recommended Actions",
        "Methodology",
        "Data Limitations",
        "Appendix: Measure Definitions",
    ):
        assert heading in html, f"missing section: {heading}"
    assert "plotly" in html.lower()  # interactive charts embedded
    assert report_data.profile.grant_name in html


def test_html_report_written_to_disk(report_data, tmp_path):
    path = write_html_report(report_data, tmp_path / "report.html")
    assert path.exists()
    assert path.stat().st_size > 10_000


def test_docx_report_written_and_valid(report_data, tmp_path):
    path = write_docx_report(report_data, tmp_path / "report.docx")
    doc = Document(str(path))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for expected in (
        "Program Overview",
        "Executive Summary",
        "Data Quality Statement",
        "Population Served",
        "Performance Measures",
        "Recommended Actions",
        "Methodology",
        "Data Limitations",
    ):
        assert any(expected in h for h in headings), f"missing heading: {expected}"
    assert doc.tables, "report should contain tables"


def test_audit_workbook_sheets_and_content(audit_flawed, prepared_flawed, tmp_path):
    path = write_audit_workbook(audit_flawed, prepared_flawed, tmp_path / "audit.xlsx")
    wb = openpyxl.load_workbook(path)
    assert set(wb.sheetnames) >= {
        "Audit Summary",
        "Issues by Rule",
        "Row-Level Issues",
        "Flagged Records",
    }
    issues_sheet = wb["Row-Level Issues"]
    assert issues_sheet.max_row == audit_flawed.total_findings + 1  # header + findings
    flagged = wb["Flagged Records"]
    header = [c.value for c in flagged[1]]
    assert "Issues Found" in header
    assert "Corrected Value(s)" in header


def test_analytics_workbook_sheets(analytics_flawed, tmp_path):
    path = write_analytics_workbook(analytics_flawed, tmp_path / "analytics.xlsx")
    wb = openpyxl.load_workbook(path)
    assert set(wb.sheetnames) >= {
        "Overview",
        "Programs",
        "Performance Measures",
        "Follow-Ups",
        "Demographics",
        "Exit Destinations",
        "Monthly Trends",
    }
    overview = wb["Overview"]
    metrics = {row[0].value: row[1].value for row in overview.iter_rows(min_row=2)}
    assert metrics["Total enrollments"] == analytics_flawed.total_enrollments


def test_concise_template_is_shorter_but_keeps_the_headline_numbers(report_data, analytics_flawed):
    full = render_html_report(report_data)
    concise = render_html_report(report_data, template="concise")
    assert len(concise) < len(full) / 2
    assert "Executive Brief" in concise
    assert "Recommended Actions" in concise
    assert "Performance Measures" in concise
    # Same source of truth: headline figures must match the full report exactly.
    assert str(analytics_flawed.total_enrollments) in concise
    assert str(analytics_flawed.total_exits) in concise
    # Detail sections belong to the full report only.
    assert "Appendix: Measure Definitions" not in concise
    assert "Demographic Summary" not in concise


def test_concise_template_written_to_disk(report_data, tmp_path):
    path = write_html_report(report_data, tmp_path / "brief.html", template="concise")
    assert path.exists()
    assert "Executive Brief" in path.read_text(encoding="utf-8")


def test_unknown_template_rejected(report_data):
    with pytest.raises(ValueError, match="Unknown report template"):
        render_html_report(report_data, template="fancy")


def test_clean_dataset_report_renders_without_audit_section(analytics_clean, profile):
    data = build_report_data(analytics_clean, None, profile)
    html = render_html_report(data)
    assert "Data Quality Statement" not in html
    assert "Executive Summary" in html


# -- The recorded trend in the funder-facing report ---------------------------
#
# The store has always held this history; until now it reached the CLI and the
# app and never the document that goes to the funder.


@pytest.fixture()
def report_with_history(tmp_path, analytics_flawed, audit_flawed, audit_clean, profile):
    from grant_assistant.history import build_history_summary, load_history, record_run
    from grant_assistant.reporting import build_report_data

    db = tmp_path / "history.db"
    record_run(profile, audit_flawed, analytics_flawed, db, label="Q1 FY26")
    summary = build_history_summary(load_history(db), audit_clean, profile.profile_id)
    return build_report_data(analytics_flawed, audit_clean, profile, history=summary)


def test_a_report_without_history_omits_the_section(analytics_flawed, audit_flawed, profile):
    """No recorded runs means no trend claim — not an empty heading."""
    from grant_assistant.reporting import build_report_data, render_html_report

    data = build_report_data(analytics_flawed, audit_flawed, profile)
    assert not data.has_history
    assert "Data Quality Over Time" not in render_html_report(data)


def test_the_html_report_states_the_movement(report_with_history):
    from grant_assistant.reporting import render_html_report

    html = render_html_report(report_with_history)
    assert "Data Quality Over Time" in html
    assert "Q1 FY26" in html
    assert report_with_history.history is not None
    assert f"{report_with_history.history.since_previous:+.1f}" in html


def test_the_word_report_states_the_movement(report_with_history, tmp_path):
    from docx import Document

    from grant_assistant.reporting import write_docx_report

    path = write_docx_report(report_with_history, tmp_path / "r.docx")
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    tables = "\n".join(
        cell.text
        for table in Document(str(path)).tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Data Quality Over Time" in text
    assert "Q1 FY26" in tables


def test_the_profile_can_turn_the_section_off(report_with_history):
    """Section selection governs the trend like every other section."""
    import dataclasses

    from grant_assistant.reporting import render_html_report

    # The profile fixture is session-scoped: editing its sections in place would
    # silently disable this section for every test that runs after this one.
    without = report_with_history.profile.model_copy(deep=True)
    without.report.sections = [s for s in without.report.sections if s != "history"]
    data = dataclasses.replace(report_with_history, profile=without)

    assert "Data Quality Over Time" not in render_html_report(data)
    assert "Data Quality Over Time" in render_html_report(report_with_history), (
        "the unmodified report must be unaffected"
    )


def test_the_report_and_the_analyst_are_given_the_same_history(
    tmp_path, analytics_flawed, audit_flawed, audit_clean, profile
):
    """One summary object, so a narrative sentence cannot contradict the table."""
    from grant_assistant.agents import DataAnalystAgent
    from grant_assistant.history import build_history_summary, load_history, record_run
    from grant_assistant.reporting import build_report_data

    db = tmp_path / "history.db"
    record_run(profile, audit_flawed, analytics_flawed, db, label="Q1")
    summary = build_history_summary(load_history(db), audit_clean, profile.profile_id)

    agent = DataAnalystAgent(analytics_flawed, audit_clean, profile, history=summary)
    data = build_report_data(analytics_flawed, audit_clean, profile, agent, summary)

    assert data.history is summary
    assert agent.fact_sheet["quality_history"]["score_change_since_previous_run"] == (
        summary.since_previous
    )


def test_the_executive_brief_states_the_movement_too(report_with_history):
    """The brief is built from the same results; silently dropping the trend hides it."""
    from grant_assistant.reporting import render_html_report

    concise = render_html_report(report_with_history, template="concise")
    assert "Data Quality Over Time" in concise
    assert report_with_history.history is not None
    assert report_with_history.history.headline() in concise


def test_the_brief_omits_the_trend_when_there_is_none(analytics_flawed, audit_flawed, profile):
    from grant_assistant.reporting import build_report_data, render_html_report

    data = build_report_data(analytics_flawed, audit_flawed, profile)
    assert "Data Quality Over Time" not in render_html_report(data, template="concise")
