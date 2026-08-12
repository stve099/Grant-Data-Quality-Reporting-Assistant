"""Chart images in the Word report.

The backend is optional, so the property that matters most is that its absence
degrades rather than breaks: a Word report without charts is still a correct
Word report. Tests therefore run both with the backend stubbed out and, when it
is genuinely installed, against real rendering.
"""

from __future__ import annotations

import pytest
from docx import Document

from grant_assistant.reporting import write_docx_report
from grant_assistant.reporting.chart_images import (
    ChartBackendError,
    chart_backend_available,
    figure_png,
    require_chart_backend,
)


@pytest.fixture(scope="module")
def report_data(analytics_flawed, audit_flawed, profile):
    from grant_assistant.reporting import build_report_data

    return build_report_data(analytics_flawed, audit_flawed, profile)


def _image_count(path) -> int:
    """Embedded images, counted from the document's related parts."""
    doc = Document(str(path))
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


@pytest.fixture(autouse=True)
def _clear_backend_cache():
    """The availability check is cached; tests must not inherit each other's."""
    chart_backend_available.cache_clear()
    yield
    chart_backend_available.cache_clear()


# -- Degradation -------------------------------------------------------------


def test_report_builds_without_a_chart_backend(report_data, tmp_path, monkeypatch):
    """The whole point: no kaleido, no charts, still a valid report."""
    monkeypatch.setattr(
        "grant_assistant.reporting.chart_images.chart_backend_available", lambda: False
    )
    path = write_docx_report(report_data, tmp_path / "no_charts.docx")
    assert path.exists()
    assert _image_count(path) == 0
    # Content is unaffected.
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "Executive Summary" in text or report_data.executive_summary[:30] in text


def test_a_failing_render_does_not_break_the_report(report_data, tmp_path, monkeypatch):
    """One bad chart must not cost the document."""
    monkeypatch.setattr(
        "grant_assistant.reporting.chart_images.chart_backend_available", lambda: True
    )

    def explode(*args, **kwargs):
        raise RuntimeError("backend crashed")

    monkeypatch.setattr("plotly.graph_objects.Figure.to_image", explode)
    path = write_docx_report(report_data, tmp_path / "broken.docx")
    assert path.exists()
    assert _image_count(path) == 0


def test_figure_png_returns_none_without_a_backend(monkeypatch):
    import plotly.graph_objects as go

    monkeypatch.setattr(
        "grant_assistant.reporting.chart_images.chart_backend_available", lambda: False
    )
    assert figure_png(go.Figure()) is None


def test_require_chart_backend_names_the_fix_when_kaleido_is_absent(monkeypatch):
    monkeypatch.setattr(
        "grant_assistant.reporting.chart_images.chart_backend_available", lambda: False
    )
    monkeypatch.setattr("grant_assistant.reporting.chart_images._kaleido_installed", lambda: False)
    with pytest.raises(ChartBackendError, match="--extra charts"):
        require_chart_backend()


def test_require_chart_backend_names_the_browser_when_that_is_what_is_missing(monkeypatch):
    """Installing the extra is useless advice when the extra is already installed."""
    monkeypatch.setattr(
        "grant_assistant.reporting.chart_images.chart_backend_available", lambda: False
    )
    monkeypatch.setattr("grant_assistant.reporting.chart_images._kaleido_installed", lambda: True)
    with pytest.raises(ChartBackendError, match="plotly_get_chrome"):
        require_chart_backend()


# -- Real rendering ----------------------------------------------------------

needs_backend = pytest.mark.skipif(
    not chart_backend_available(), reason="the charts extra is not installed"
)


@needs_backend
def test_charts_are_embedded_when_the_backend_is_available(report_data, tmp_path):
    path = write_docx_report(report_data, tmp_path / "with_charts.docx")
    assert _image_count(path) > 0


@needs_backend
def test_rendered_png_is_a_png():
    import plotly.graph_objects as go

    png = figure_png(go.Figure(go.Bar(x=["a"], y=[1])), width=400, height=250)
    assert png is not None
    assert png.startswith(b"\x89PNG\r\n\x1a\n")
