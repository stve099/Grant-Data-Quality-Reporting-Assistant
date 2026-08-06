"""Generate a professional Microsoft Word grant report with python-docx."""

from __future__ import annotations

import io
import logging
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from grant_assistant.reporting.chart_images import CHART_WIDTH_INCHES, figure_png
from grant_assistant.reporting.context import ReportData
from grant_assistant.reporting.formatting import format_value as _fmt

logger = logging.getLogger(__name__)

# Tokens from docs/design_system.md (brand-deep blue, secondary ink).
BRAND = RGBColor(0x1C, 0x5C, 0xAB)
MUTED = RGBColor(0x52, 0x51, 0x4E)


def _add_table(doc: DocumentType, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def _bullets(doc: DocumentType, items: list[str]) -> None:
    if not items:
        doc.add_paragraph("None noted.", style="List Bullet")
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_chart(doc: DocumentType, charts: dict[str, object], key: str, caption: str = "") -> bool:
    """Insert a chart if it exists and can be rendered. Returns whether it was.

    Every step is optional: the chart may not apply to this dataset, and the
    static backend may not be installed. A Word report without charts is still a
    correct Word report, so absence never interrupts the build.
    """
    figure = charts.get(key)
    if figure is None:
        return False
    png = figure_png(figure)  # type: ignore[arg-type]
    if png is None:
        return False
    doc.add_picture(io.BytesIO(png), width=Inches(CHART_WIDTH_INCHES))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = MUTED
    return True


def write_docx_report(report: ReportData, path: str | Path) -> Path:
    """Build the Word report and write it to ``path``."""
    from grant_assistant.analytics.charts import standard_chart_set

    doc = Document()
    a = report.analytics
    charts: dict[str, object] = standard_chart_set(a, report.audit)

    # --- Cover page ---------------------------------------------------------
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.title)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report.profile.grant_name)
    run.font.size = Pt(18)
    run.font.color.rgb = MUTED

    for text in (
        f"Reporting Period: {report.period_label}",
        f"Prepared by: {report.profile.report.prepared_by}",
        f"Generated: {report.generated_at:%B %d, %Y}",
        "Synthetic demonstration data — contains no real client information.",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = MUTED
    doc.add_page_break()

    # --- Program overview ---------------------------------------------------
    doc.add_heading("Program Overview", level=1)
    doc.add_paragraph(
        report.profile.description
        or "This report summarizes enrollment, outcomes, income, follow-up, and data "
        "quality results for the funded programs during the reporting period."
    )
    _add_table(
        doc,
        ["Program", "Description"],
        [[p.name, p.description or "—"] for p in report.profile.programs],
    )

    # --- Executive summary --------------------------------------------------
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.executive_summary)
    mode = (
        "AI-assisted narrative grounded in deterministic calculations."
        if report.ai_generated_narrative
        else "Deterministic narrative (non-AI mode)."
    )
    p = doc.add_paragraph()
    r = p.add_run(f"Narrative mode: {mode}")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    # --- Data quality -------------------------------------------------------
    if report.audit is not None:
        doc.add_heading("Data Quality Statement", level=1)
        doc.add_paragraph(report.audit.executive_summary())
        _add_table(
            doc,
            ["Severity", "Findings"],
            [
                [sev.title() if sev != "info" else "Informational", str(count)]
                for sev, count in report.audit.issue_count_by_severity.items()
            ],
        )
        if report.audit.blocking_issues:
            doc.add_paragraph("Blocking issues that must be resolved before submission:")
            _bullets(
                doc,
                [
                    f"{i.rule_name} — {i.record_count} record(s)"
                    for i in report.audit.blocking_issues
                ],
            )

    # --- Population served --------------------------------------------------
    doc.add_heading("Population Served", level=1)
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total enrollments", str(a.total_enrollments)],
            ["Households served", str(a.households_served)],
            ["Total individuals", str(a.total_individuals)],
            ["Adults", str(a.total_adults)],
            ["Children", str(a.total_children)],
            ["Active enrollments", str(a.active_enrollments)],
        ],
    )

    # --- Demographics -------------------------------------------------------
    doc.add_heading("Demographic Summary", level=1)
    for label, rows in report.demographic_tables():
        doc.add_heading(label, level=2)
        _add_table(doc, [label, "Clients"], [[value, str(count)] for value, count in rows])

    # --- Enrollment & exits -------------------------------------------------
    doc.add_heading("Enrollment & Exit Metrics", level=1)
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Total exits", str(a.total_exits)],
            ["Exit rate", _fmt(a.exit_rate, "percent")],
            ["Successful exits", str(a.successful_exits)],
            ["Successful exit rate", _fmt(a.successful_exit_rate, "percent")],
            ["Permanent housing exits", str(a.permanent_housing_exits)],
            ["Permanent housing rate", _fmt(a.permanent_housing_rate, "percent")],
        ],
    )

    _add_chart(doc, charts, "enrollment_trends", "Monthly enrollments and exits")

    # --- Housing outcomes ---------------------------------------------------
    doc.add_heading("Housing Outcomes", level=1)
    _add_table(
        doc,
        ["Exit destination", "Exits"],
        [[dest, str(count)] for dest, count in a.exit_destination_breakdown.items()],
    )

    _add_chart(doc, charts, "exit_destinations", "Where households went at exit")

    # --- Income outcomes ----------------------------------------------------
    doc.add_heading("Income Outcomes", level=1)
    _add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Average entry income", _fmt(a.avg_entry_income, "currency")],
            ["Average exit income", _fmt(a.avg_exit_income, "currency")],
            ["Average income change", _fmt(a.avg_income_change, "currency")],
            ["Median income change", _fmt(a.median_income_change, "currency")],
            ["% households increasing income", _fmt(a.pct_income_increased, "percent")],
            ["Exits with complete income data", str(a.n_income_pairs)],
        ],
    )

    _add_chart(doc, charts, "income_change", "Income change from entry to exit")

    # --- Follow-ups ---------------------------------------------------------
    doc.add_heading("Follow-Up Outcomes", level=1)
    _add_table(
        doc,
        ["Milestone", "Due", "Completed", "Overdue", "Completion rate"],
        [
            [
                f.label,
                str(f.due),
                str(f.completed_of_due),
                str(f.overdue),
                _fmt(f.completion_rate, "percent"),
            ]
            for f in a.followups
        ],
    )

    _add_chart(doc, charts, "followups", "Follow-up completion by interval")

    # --- Performance measures -----------------------------------------------
    doc.add_heading("Performance Measures", level=1)
    _add_table(
        doc,
        ["Measure", "Target", "Actual", "Status"],
        [
            [
                m.name + (" (small sample)" if m.small_sample else ""),
                _fmt(m.target, m.unit),
                _fmt(m.actual, m.unit),
                "Met" if m.met else ("Not met" if m.met is False else "No data"),
            ]
            for m in a.measures
        ],
    )

    _add_chart(doc, charts, "goal_vs_actual", "Performance measures against target")

    # --- Program comparison -------------------------------------------------
    doc.add_heading("Program Comparison", level=1)
    _add_table(
        doc,
        [
            "Program",
            "Enrollments",
            "Exits",
            "Successful exit rate",
            "Permanent housing rate",
            "Avg income change",
        ],
        [
            [
                p.program,
                str(p.enrollments),
                str(p.exits),
                _fmt(p.successful_exit_rate, "percent"),
                _fmt(p.permanent_housing_rate, "percent"),
                _fmt(p.avg_income_change, "currency"),
            ]
            for p in a.programs
        ],
    )

    _add_chart(doc, charts, "program_comparison", "Enrollments and exits by program")
    _add_chart(doc, charts, "outcome_rates", "Outcome rates by program")

    # --- Findings and recommendations ---------------------------------------
    doc.add_heading("Key Findings", level=1)
    _bullets(doc, report.insights.key_findings + report.insights.notable_trends)
    doc.add_heading("Challenges & Risks", level=1)
    _bullets(
        doc,
        report.insights.program_concerns
        + report.insights.data_quality_risks
        + report.insights.anomalies,
    )
    doc.add_heading("Recommended Actions", level=1)
    _bullets(doc, report.insights.recommended_actions)

    # --- Methodology and limitations ----------------------------------------
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "All metrics are calculated deterministically from the uploaded data extract using "
        f"the '{report.profile.profile_id}' grant profile. Field mappings translate source "
        "columns onto a canonical schema; program aliases are normalized; exact duplicate "
        "enrollments are removed before analysis; outcome categories and successful-exit "
        "definitions come from the profile's exit destination mappings; follow-up due dates "
        "are derived from exit dates plus the configured schedule. AI narrative (when "
        "enabled) is generated strictly from these pre-calculated metrics."
    )
    doc.add_heading("Data Limitations", level=1)
    _bullets(doc, report.data_limitations())

    # --- Appendix -----------------------------------------------------------
    doc.add_heading("Appendix: Measure Definitions", level=1)
    _add_table(
        doc,
        ["Measure", "Definition"],
        [[name, definition] for name, definition in report.measure_definitions()],
    )

    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    logger.info("Wrote Word report to %s", path)
    return path
